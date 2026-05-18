import argparse
import re

from docx import Document
from docx.shared import Pt


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
ORDERED_RE = re.compile(r"^(\d+)\.\s+(.*)$")


def _add_code_block(document: Document, code_lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run("\n".join(code_lines))
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def _parse_table(lines: list[str], start: int) -> tuple[int, list[list[str]]]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and "|" in lines[i]:
        raw = lines[i].strip()
        if not raw:
            break
        rows.append([cell.strip() for cell in raw.strip("|").split("|")])
        i += 1

    if len(rows) >= 2:
        sep = rows[1]
        if all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in sep):
            rows = [rows[0]] + rows[2:]
        else:
            return start, []
    else:
        return start, []

    return i, rows


def _add_table(document: Document, rows: list[list[str]]) -> None:
    col_count = max(len(r) for r in rows) if rows else 0
    if col_count <= 0:
        return
    table = document.add_table(rows=len(rows), cols=col_count)
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            value = row[c_idx] if c_idx < len(row) else ""
            table.cell(r_idx, c_idx).text = value


def convert(markdown_text: str) -> Document:
    document = Document()

    lines = markdown_text.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        text = "\n".join([line.rstrip() for line in paragraph_lines]).strip()
        paragraph_lines = []
        if text:
            document.add_paragraph(text)

    while i < len(lines):
        line = lines[i].rstrip("\n")

        if line.strip().startswith("```"):
            if in_code:
                _add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                flush_paragraph()
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            flush_paragraph()
            i += 1
            continue

        heading_match = HEADING_RE.match(line.strip())
        if heading_match:
            flush_paragraph()
            level = min(len(heading_match.group(1)), 6)
            title = heading_match.group(2).strip()
            document.add_heading(title, level=level)
            i += 1
            continue

        ordered_match = ORDERED_RE.match(line.strip())
        if ordered_match:
            flush_paragraph()
            document.add_paragraph(ordered_match.group(2).strip(), style="List Number")
            i += 1
            continue

        if line.lstrip().startswith(("-", "*")) and line.lstrip()[1:2] == " ":
            flush_paragraph()
            document.add_paragraph(line.lstrip()[2:].strip(), style="List Bullet")
            i += 1
            continue

        if "|" in line:
            flush_paragraph()
            next_i, rows = _parse_table(lines, i)
            if rows:
                _add_table(document, rows)
                i = next_i
                continue

        paragraph_lines.append(line)
        i += 1

    flush_paragraph()

    if in_code and code_lines:
        _add_code_block(document, code_lines)

    return document


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input_md")
    parser.add_argument("output_docx")
    args = parser.parse_args()

    with open(args.input_md, "r", encoding="utf-8") as handle:
        markdown_text = handle.read()

    document = convert(markdown_text)
    document.save(args.output_docx)


if __name__ == "__main__":
    main()

