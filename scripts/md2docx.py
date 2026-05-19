import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys

INPUT_FILE = r"D:\MyProjects\bishe-finnal\docs\thesis\本科毕业论文-基于用户行为日志的电商智能分析与决策支持系统设计与实现.md"
OUTPUT_FILE = r"D:\MyProjects\bishe-finnal\docs\thesis\本科毕业论文-电商智能分析与决策支持系统.docx"

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

def set_font(run, name='宋体', size=Pt(12), bold=False, color=None):
    run.font.name = name
    run.font.size = size
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = color

def add_heading_styled(text, level):
    p = doc.add_paragraph()
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_font(run, '黑体', Pt(22), bold=True)
    elif level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_font(run, '黑体', Pt(16), bold=True)
    elif level == 2:
        run = p.add_run(text)
        set_font(run, '黑体', Pt(14), bold=True)
    elif level == 3:
        run = p.add_run(text)
        set_font(run, '黑体', Pt(12), bold=True)
    else:
        run = p.add_run(text)
        set_font(run, '宋体', Pt(12), bold=True)
    return p

def add_paragraph_styled(text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    process_inline(p, text)
    return p

def process_inline(p, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_font(run, '宋体', Pt(12), bold=True)
        else:
            cleaned = part.replace('\\*', '*')
            if cleaned:
                run = p.add_run(cleaned)
                set_font(run, '宋体', Pt(12))

def add_table_with_data(headers, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h.strip())
        set_font(run, '宋体', Pt(10), bold=True)

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text.strip())
            set_font(run, '宋体', Pt(10))

    doc.add_paragraph()
    return table

def parse_table(lines, start_idx):
    header_line = lines[start_idx].strip()
    if '|' not in header_line:
        return None, start_idx

    headers = [c.strip() for c in header_line.split('|') if c.strip()]

    sep_idx = start_idx + 1
    if sep_idx < len(lines) and '---' in lines[sep_idx]:
        pass
    else:
        return None, start_idx

    rows = []
    idx = sep_idx + 1
    while idx < len(lines):
        line = lines[idx].strip()
        if '|' in line and not line.startswith('#'):
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if len(cells) >= len(headers):
                rows.append(cells[:len(headers)])
                idx += 1
                continue
        break

    return (headers, rows), idx

def parse_markdown(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        heading_match = re.match(r'^(#{1,6})\s+(.+)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            add_heading_styled(text, level)
            i += 1
            continue

        table_result = parse_table(lines, i)
        if table_result[0] is not None:
            headers, rows = table_result[0]
            add_table_with_data(headers, rows)
            i = table_result[1]
            continue

        if line.strip().startswith('【'):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(24)
            run = p.add_run(line.strip())
            set_font(run, '宋体', Pt(10.5), bold=False)
            i += 1
            continue

        list_match = re.match(r'^(\s*[-*]\s+|\s*\(\d+\)\s*|\s*（\d+）\s*|\s*\d+[\.\、]\s*)(.+)', line)
        if list_match:
            text = list_match.group(2).strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            process_inline(p, text)
            i += 1
            continue

        if re.match(r'^图\d+[\-\u2014\u2015].*', line.strip()):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.strip())
            set_font(run, '宋体', Pt(10.5))
            i += 1
            continue

        if re.match(r'^表\d+[\-\u2014\u2015].*', line.strip()):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.strip())
            set_font(run, '宋体', Pt(10.5), bold=True)
            i += 1
            continue

        add_paragraph_styled(line.strip())
        i += 1

parse_markdown(INPUT_FILE)

doc.save(OUTPUT_FILE)
print(f"转换完成！\n输出文件: {OUTPUT_FILE}")